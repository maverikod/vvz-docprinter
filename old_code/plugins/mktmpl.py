import sys
import os
import warnings

from plugins.utils import plugin_result
from plugins.utils import PluginError

from docx import Document
from docxtpl import DocxTemplate
import re

def replace_text_with_pattern(text, pattern, replacement):
    """
    Заменяет все подстроки в тексте, удовлетворяющие заданному шаблону,
    на указанную замену.
    """
    return re.sub(pattern, replacement, text)

def replace_run(run, pattern):
    """
    Заменяет текст внутри run, удовлетворяющий заданному шаблону,
    на новый формат.
    """
    #if re.search(pattern, run.text):
    #    updated_text = replace_text_with_pattern(run.text, pattern, r'{%tr for \1Str in \1 %}\2 {{\1Str.\3}}\4 {{\1Str.\5}}\6 {{\1Str.\7}}{%tr endfor %}')
    #    run.text = updated_text

    # Замена символа [ на {{
    updated_text = replace_text_with_pattern(run.text, r'\[', r'{{')
    run.text = updated_text

    # Замена символа ] на }}
    updated_text = replace_text_with_pattern(run.text, r'\]', r'}}')
    run.text = updated_text

def replace_text(doc):
    # Шаблон для поиска подстрок
    pattern = r'{\s*([\w_]+)>\s*([^[\]]+)\s+\[([\w_]+)\]\s*([^[\]]+)\s+\[([\w_]+)\]\s*([^[\]]+)\s+\[([\w_]+)\]\s*([^[\]]+)\s*<}'

    # Замена подстрок на новый формат
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            replace_run(run, pattern)

    # Обработка текста внутри таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        replace_run(run, pattern)

def process_docx(file_path, dest_directory):
    try:
        doc = Document(file_path)
        replace_text(doc)

        # Определение нового имени файла
        file_name = os.path.basename(file_path)
        new_file_path = os.path.join(dest_directory, file_name)

        doc.save(new_file_path)
        print(f"Создан файл '{new_file_path}'")
        return plugin_result(status=0)
    except PermissionError:
        print(f"Ошибка: Нет разрешения для записи в файл '{new_file_path}'")
        return plugin_result(status=971,data={"filename":new_file_path})

def run(data,locale,msg):
    if not isinstance(data,dict):
       return plugin_result(status=992)

    if not "indir" in data or not "outdir" in data:
        return plugin_result(973)
    
    if not os.path.isdir(data["indir"]):
        return plugin_result(status=1,data={"indir":data["indir"]})

    if os.path.isfile(data["outdir"]):
        return plugin_result(status=972,data={"outdir":data["outdir"]})
    elif not isinstance(data["outdir"],str):
        return plugin_result(status=971,data={"outdir":data["outdir"]})

    input_directory  = data["indir"]
    output_directory = data["outdir"]

    # Проверка наличия каталога назначения и его создание, если не существует
    if not os.path.isdir(output_directory):
        try:
            os.makedirs(output_directory)
            print(f"Каталог назначения '{output_directory}' успешно создан.")
        except OSError:
            print(f"Ошибка: Не удалось создать каталог назначения '{output_directory}'")
            return plugin_result(status=970,data={"outdir":data["outdir"]})

    try:
        # Проверка наличия входного каталога
        if not os.path.isdir(input_directory):
            return plugin_result(status=969,data={"indir":data["indir"]})

        # Поиск файлов с расширением .docx в указанном каталоге
        for file_name in os.listdir(input_directory):
            if file_name.lower().endswith(".docx"):
                file_path = os.path.join(input_directory, file_name)
                process_docx(file_path, output_directory)
        return plugin_result(status=0)

    except FileNotFoundError as e:
        print(f"Ошибка: {str(e)}")
        return plugin_result(status=1,data={"file":file_path})
    except PermissionError as e:
        return plugin_result(status=4,data={"file":file_path})
