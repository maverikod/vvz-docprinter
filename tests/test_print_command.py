"""
Tests for ``PrintCommand`` (base64 ZIP contract).

Author: Vasiliy Zdanovskiy
email: vasilyvz@gmail.com
"""

from __future__ import annotations

import asyncio
import base64
import json
import uuid
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Iterator

import pytest
from docx import Document
from mcp_proxy_adapter.config import get_config

from docprinter.commands import print_command as print_command_module
from docprinter.commands.print_command import PrintCommand

REPO_ROOT = Path(__file__).resolve().parents[1]


@pytest.fixture
def docprinter_env(tmp_path: Path) -> tuple[Path, Path]:
    """Route docprinter output/work under tmp_path for isolation."""
    out = tmp_path / "out"
    wk = tmp_path / "wk"
    cfg = get_config()
    cfg.config_data["docprinter"] = {
        "output_dir": str(out),
        "work_dir": str(wk),
        "output_ttl_seconds": 3600,
        "sweep_interval_seconds": 300,
    }
    return out, wk


def _minimal_docx_bytes(body: str) -> bytes:
    doc = Document()
    doc.add_paragraph(body)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _zip_archive_bytes(data_obj: object, template_body: str) -> bytes:
    zbuf = BytesIO()
    with zipfile.ZipFile(zbuf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("data.json", json.dumps(data_obj))
        zf.writestr("template.docx", _minimal_docx_bytes(template_body))
    return zbuf.getvalue()


def test_print_success_returns_result_id(
    docprinter_env: tuple[Path, Path], monkeypatch: pytest.MonkeyPatch
) -> None:
    """Valid archive yields result_id, output file exists, work dir removed."""
    out, wk = docprinter_env
    u_request = uuid.UUID("11111111-1111-4111-8111-111111111111")
    u_result = uuid.UUID("22222222-2222-4222-8222-222222222222")
    seq: Iterator[uuid.UUID] = iter([u_request, u_result])

    def _fake_uuid4() -> uuid.UUID:
        return next(seq)

    monkeypatch.setattr(print_command_module.uuid, "uuid4", _fake_uuid4)

    raw = _zip_archive_bytes({"name": "Ada"}, "Hello {{ name }}")
    archive_b64 = base64.b64encode(raw).decode("ascii")

    async def _run() -> dict:
        cmd = PrintCommand()
        r = await cmd.execute(archive=archive_b64)
        return r.to_dict()

    d = asyncio.run(_run())
    assert d["success"] is True
    rid = d["data"]["result_id"]
    uuid.UUID(rid)
    assert rid == str(u_result)
    assert (out / f"{rid}.docx").is_file()
    assert not (wk / str(u_request)).exists()


def test_print_invalid_base64(docprinter_env: tuple[Path, Path]) -> None:
    """Non-base64 payload yields INVALID_BASE64."""

    async def _run() -> dict:
        cmd = PrintCommand()
        r = await cmd.execute(archive="!!!not-base64!!!")
        return r.to_dict()

    d = asyncio.run(_run())
    assert d["success"] is False
    assert d["error"]["data"]["error_code"] == "INVALID_BASE64"


def test_print_invalid_archive(docprinter_env: tuple[Path, Path]) -> None:
    """Non-ZIP bytes yield INVALID_ARCHIVE."""
    raw = b"not a zip"
    archive_b64 = base64.b64encode(raw).decode("ascii")

    async def _run() -> dict:
        cmd = PrintCommand()
        r = await cmd.execute(archive=archive_b64)
        return r.to_dict()

    d = asyncio.run(_run())
    assert d["success"] is False
    assert d["error"]["data"]["error_code"] == "INVALID_ARCHIVE"


def test_print_archive_missing_files(docprinter_env: tuple[Path, Path]) -> None:
    """ZIP missing template.docx reports ARCHIVE_MISSING_REQUIRED_FILES."""
    zbuf = BytesIO()
    with zipfile.ZipFile(zbuf, "w") as zf:
        zf.writestr("data.json", "{}")
    archive_b64 = base64.b64encode(zbuf.getvalue()).decode("ascii")

    async def _run() -> dict:
        cmd = PrintCommand()
        r = await cmd.execute(archive=archive_b64)
        return r.to_dict()

    d = asyncio.run(_run())
    assert d["success"] is False
    assert d["error"]["data"]["error_code"] == "ARCHIVE_MISSING_REQUIRED_FILES"
    assert "template.docx" in json.dumps(d["error"]["data"])


def test_print_unsafe_archive_member(
    tmp_path: Path, docprinter_env: tuple[Path, Path]
) -> None:
    """Path members inside ZIP are rejected; nothing escapes work_dir."""
    _, _ = docprinter_env
    zbuf = BytesIO()
    with zipfile.ZipFile(zbuf, "w") as zf:
        zf.writestr("../evil.txt", b"x")
    archive_b64 = base64.b64encode(zbuf.getvalue()).decode("ascii")

    async def _run() -> dict:
        cmd = PrintCommand()
        r = await cmd.execute(archive=archive_b64)
        return r.to_dict()

    d = asyncio.run(_run())
    assert d["success"] is False
    assert d["error"]["data"]["error_code"] == "UNSAFE_ARCHIVE_MEMBER"
    assert not list(tmp_path.rglob("evil.txt"))


def test_print_invalid_data_json(docprinter_env: tuple[Path, Path]) -> None:
    """Non-object JSON root yields INVALID_DATA_TYPE."""
    zbuf = BytesIO()
    with zipfile.ZipFile(zbuf, "w") as zf:
        zf.writestr("data.json", json.dumps("just a string"))
        zf.writestr("template.docx", _minimal_docx_bytes("x"))
    archive_b64 = base64.b64encode(zbuf.getvalue()).decode("ascii")

    async def _run() -> dict:
        cmd = PrintCommand()
        r = await cmd.execute(archive=archive_b64)
        return r.to_dict()

    d = asyncio.run(_run())
    assert d["success"] is False
    assert d["error"]["data"]["error_code"] == "INVALID_DATA_TYPE"


def test_print_legacy_plugin_result_wrapper_is_unwrapped(
    docprinter_env: tuple[Path, Path], monkeypatch: pytest.MonkeyPatch
) -> None:
    """read_json-style {exitcode, data: <full job>} still yields inner Jinja dict."""
    out, _wk = docprinter_env
    u_request = uuid.UUID("cccccccc-cccc-4ccc-8ccc-ccccccccccc0")
    u_result = uuid.UUID("cccccccc-cccc-4ccc-8ccc-ccccccccccc1")
    seq: Iterator[uuid.UUID] = iter([u_request, u_result])

    def _fake_uuid4() -> uuid.UUID:
        return next(seq)

    monkeypatch.setattr(print_command_module.uuid, "uuid4", _fake_uuid4)

    job = {
        "cmd": "docbytpl",
        "data": {
            "template": "/tmp/ignored.docx",
            "outfile": "/tmp/ignored_out.docx",
            "data": {"marker": "WRAPPER_OK"},
        },
    }
    wrapped = {"exitcode": 0, "data": job, "errstr": ""}
    zbuf = BytesIO()
    with zipfile.ZipFile(zbuf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("data.json", json.dumps(wrapped))
        zf.writestr("template.docx", _minimal_docx_bytes("Marker: {{ marker }}"))
    archive_b64 = base64.b64encode(zbuf.getvalue()).decode("ascii")

    async def _run() -> None:
        cmd = PrintCommand()
        r = await cmd.execute(archive=archive_b64)
        assert r.to_dict()["success"] is True

    asyncio.run(_run())
    out_doc = out / f"{u_result}.docx"
    assert out_doc.is_file()
    doc = Document(str(out_doc))
    joined = "\n".join(p.text for p in doc.paragraphs)
    assert "WRAPPER_OK" in joined


def test_print_v8_uni_fixture_not_empty_wrong_context(
    docprinter_env: tuple[Path, Path],
) -> None:
    """
    Regression: rendering the whole job JSON (wrong context) yields ~18 420 bytes
    for ``v8_GOPJOV_5a`` + UNI act template; correct unwrap fills ``СуммаПрописьюАнгл``.
    """
    data_path = REPO_ROOT / "test-data" / "v8_GOPJOV_5a.json"
    tpl_path = REPO_ROOT / "test-data" / "UNI Композитор (Composer) МАКЕТ АКТА.docx"
    if not data_path.is_file() or not tpl_path.is_file():
        pytest.skip("test-data fixtures not present")

    out, _wk = docprinter_env
    zbuf = BytesIO()
    with zipfile.ZipFile(zbuf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("data.json", data_path.read_bytes())
        zf.writestr("template.docx", tpl_path.read_bytes())
    archive_b64 = base64.b64encode(zbuf.getvalue()).decode("ascii")

    async def _run() -> Path:
        cmd = PrintCommand()
        r = await cmd.execute(archive=archive_b64)
        assert r.to_dict()["success"] is True
        rid = r.to_dict()["data"]["result_id"]
        return out / f"{rid}.docx"

    out_doc = asyncio.run(_run())
    blob = out_doc.read_bytes()
    assert len(blob) != 18420, (
        "output size 18420 means whole JSON was passed to Jinja "
        "(legacy bug / wrong print_command import)"
    )
    with zipfile.ZipFile(BytesIO(blob)) as zdoc:
        xml = zdoc.read("word/document.xml").decode("utf-8", errors="replace")
    assert "Seven hundred" in xml


def test_print_flat_docbytpl_triple_at_json_root(
    docprinter_env: tuple[Path, Path], monkeypatch: pytest.MonkeyPatch
) -> None:
    """``data.json`` may be only the docbytpl triple (what starter passes to plugin)."""
    out, _wk = docprinter_env
    u_request = uuid.UUID("dddddddd-dddd-4ddd-bddd-dddddddddd00")
    u_result = uuid.UUID("dddddddd-dddd-4ddd-bddd-dddddddddd01")
    seq: Iterator[uuid.UUID] = iter([u_request, u_result])

    def _fake_uuid4() -> uuid.UUID:
        return next(seq)

    monkeypatch.setattr(print_command_module.uuid, "uuid4", _fake_uuid4)

    triple_only = {
        "template": "/tmp/ignored.docx",
        "outfile": "/tmp/ignored_out.docx",
        "data": {"marker": "TRIPLE_ROOT_OK"},
    }
    zbuf = BytesIO()
    with zipfile.ZipFile(zbuf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("data.json", json.dumps(triple_only))
        zf.writestr("template.docx", _minimal_docx_bytes("M: {{ marker }}"))
    archive_b64 = base64.b64encode(zbuf.getvalue()).decode("ascii")

    async def _run() -> None:
        cmd = PrintCommand()
        r = await cmd.execute(archive=archive_b64)
        assert r.to_dict()["success"] is True

    asyncio.run(_run())
    out_doc = out / f"{u_result}.docx"
    assert out_doc.is_file()
    doc = Document(str(out_doc))
    joined = "\n".join(p.text for p in doc.paragraphs)
    assert "TRIPLE_ROOT_OK" in joined


def test_print_legacy_docbytpl_nested_data_is_jinja_context(
    docprinter_env: tuple[Path, Path], monkeypatch: pytest.MonkeyPatch
) -> None:
    """Legacy job JSON: Jinja sees ``root['data']['data']``, not the envelope."""
    out, _wk = docprinter_env
    u_request = uuid.UUID("bbbbbbbb-bbbb-4bbb-8bbb-bbbbbbbbbbb0")
    u_result = uuid.UUID("bbbbbbbb-bbbb-4bbb-8bbb-bbbbbbbbbbb1")
    seq: Iterator[uuid.UUID] = iter([u_request, u_result])

    def _fake_uuid4() -> uuid.UUID:
        return next(seq)

    monkeypatch.setattr(print_command_module.uuid, "uuid4", _fake_uuid4)

    legacy_root = {
        "cmd": "docbytpl",
        "data": {
            "template": "/tmp/ignored.docx",
            "outfile": "/tmp/ignored_out.docx",
            "data": {"marker": "LEGACY_OK"},
        },
    }
    zbuf = BytesIO()
    with zipfile.ZipFile(zbuf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("data.json", json.dumps(legacy_root))
        zf.writestr("template.docx", _minimal_docx_bytes("Marker: {{ marker }}"))
    archive_b64 = base64.b64encode(zbuf.getvalue()).decode("ascii")

    async def _run() -> None:
        cmd = PrintCommand()
        r = await cmd.execute(archive=archive_b64)
        assert r.to_dict()["success"] is True

    asyncio.run(_run())
    out_doc = out / f"{u_result}.docx"
    assert out_doc.is_file()
    doc = Document(str(out_doc))
    joined = "\n".join(p.text for p in doc.paragraphs)
    assert "LEGACY_OK" in joined


def test_print_template_undefined_var(docprinter_env: tuple[Path, Path]) -> None:
    """Undefined Jinja variable maps to TEMPLATE_UNDEFINED."""
    raw = _zip_archive_bytes({}, "Value: {{ missing_var.foo }}")
    archive_b64 = base64.b64encode(raw).decode("ascii")

    async def _run() -> dict:
        cmd = PrintCommand()
        r = await cmd.execute(archive=archive_b64)
        return r.to_dict()

    d = asyncio.run(_run())
    assert d["success"] is False
    assert d["error"]["data"]["error_code"] == "TEMPLATE_UNDEFINED"
    assert d["error"]["code"] == -32603


def test_print_schema_required_fields() -> None:
    schema = PrintCommand.get_schema()
    assert schema["required"] == ["archive"]
    assert schema.get("additionalProperties") is False


def test_print_result_schema() -> None:
    s = PrintCommand.get_result_schema()
    assert s.get("title") == "print_success"
    assert "result_id" in s["properties"]["data"]["required"]


def test_work_dir_cleaned_on_error(
    docprinter_env: tuple[Path, Path], monkeypatch: pytest.MonkeyPatch
) -> None:
    """Work directory for the request is removed after failed execute."""
    _, wk = docprinter_env
    u_request = uuid.UUID("aaaaaaaa-aaaa-4aaa-8aaa-aaaaaaaaaaaa")

    monkeypatch.setattr(print_command_module.uuid, "uuid4", lambda: u_request)

    raw = b"not a zip"
    archive_b64 = base64.b64encode(raw).decode("ascii")

    async def _run() -> None:
        cmd = PrintCommand()
        await cmd.execute(archive=archive_b64)

    asyncio.run(_run())
    assert not (wk / str(u_request)).exists()
